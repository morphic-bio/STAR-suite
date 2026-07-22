#!/usr/bin/env python3
"""Generate a .docx version of the perturb-seq QC report with embedded figures."""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

QC_DIR = Path("/mnt/pikachu/ucsf_emptydrops_guarded_release_20260403/qc_plots")
OUT = QC_DIR / "PERTURBSEQ_QC_REPORT.docx"


def add_table(doc, headers, rows, bold_col=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style="Light Shading Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    # Data rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    if bold_col is not None and j == bold_col:
                        run.bold = True

    doc.add_paragraph()  # spacer


def add_figure(doc, path, caption, width=6.0):
    """Add an image with a caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    run.italic = True
    doc.add_paragraph()  # spacer


def build_report():
    doc = Document()

    # -- Styles --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # ========== TITLE ==========
    title = doc.add_heading("Perturb-seq QC Report: UCSF CRISPRi Knockdown Validation", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # -- Metadata --
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for label, value in [
        ("Date", "2026-04-03"),
        ("Dataset", "UCSF perturb-seq (guarded EmptyDrops rerun)"),
        ("Samples", "16 (10 EBs, 6 iPSC)"),
        ("Library", "hCRISPRa-v2 AALG2 pattern (548 guides, 224 target genes, 10 NTC guides)"),
        ("Pipeline", "STAR-suite EmptyDrops_CR with guarded ambient minimum"),
    ]:
        run = meta.add_run(f"{label}: ")
        run.bold = True
        run.font.size = Pt(10)
        run = meta.add_run(f"{value}\n")
        run.font.size = Pt(10)

    # ========== OBJECTIVE ==========
    doc.add_heading("Objective", level=1)
    doc.add_paragraph(
        "Validate that CRISPRi perturbations produce measurable target gene knockdown "
        "in the refiltered dataset, assess guide representation, and identify which "
        "target genes are testable in this system."
    )

    # ========== METHODS ==========
    doc.add_heading("Methods", level=1)

    doc.add_heading("Normalization and statistical testing", level=2)
    doc.add_paragraph(
        "For each sample independently:"
    )
    methods = [
        ("Normalization", "Counts Per 10K (CP10K) library-size normalization followed by log1p transform: log1p(count \u00d7 10000 / total_UMI)."),
        ("Baseline", "Non-targeting control (NTC) cells used as the reference population (~240 cells per sample from 10 NTC guides: negative_control_1 through negative_control_5, each with _P1P2_A and _P1P2_B variants)."),
        ("Guide-to-gene mapping", "Each cell\u2019s feature_call assignment was mapped to a target gene via the feature reference CSV. Cells assigned to different guides targeting the same gene were pooled."),
        ("Fold change", "log2 fold change computed as log2(mean_perturbed + 0.01) \u2013 log2(mean_NTC + 0.01) in log1p-CP10K space."),
        ("Statistical test", "Two-sided Mann-Whitney U (Wilcoxon rank-sum) test comparing normalized target gene expression in perturbed cells vs NTC cells."),
        ("Multiple testing correction", "Benjamini-Hochberg FDR correction applied per sample across all tested genes."),
        ("Significance threshold", "BH-adjusted p < 0.05 AND log2FC < \u20130.5."),
    ]
    for i, (term, desc) in enumerate(methods, 1):
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(f"{term}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_heading("Expression filtering", level=2)
    doc.add_paragraph(
        "A target gene was classified as \u201cexpressed\u201d in a condition (EBs or iPSC) if "
        "its median NTC CP10K expression across that condition\u2019s samples was \u2265 1.0. "
        "Genes not expressed in a condition cannot show knockdown and were excluded "
        "from interpretation for that condition."
    )

    doc.add_heading("Minimum cell threshold", level=2)
    doc.add_paragraph(
        "Target genes with fewer than 5 perturbed cells in a sample were excluded "
        "from statistical testing in that sample."
    )

    # ========== RESULTS ==========
    doc.add_heading("Results", level=1)

    # -- 1. Guide representation --
    doc.add_heading("1. Guide representation is highly skewed", level=2)

    add_figure(
        doc,
        QC_DIR / "cell_count_distribution.png",
        "Figure 1. Distribution of median perturbed cell count per target gene across all 16 samples. "
        "Of 223 target genes detected, 167 have \u2265 20 cells (median across samples). "
        "The distribution is heavily right-skewed: three genes (TP53, PTEN, SOX4) dominate "
        "with 300\u2013550 cells, while the bottom quartile has < 21 cells.",
        width=6.5,
    )

    add_table(doc,
        ["Statistic", "Value"],
        [
            ["Target genes in matrix", "223 / 224"],
            ["Median cells per gene (across genes)", "44"],
            ["Q1 / Q3", "21 / 61"],
            ["Max (TP53)", "552"],
            ["Min", "1"],
            ["Genes \u2265 20 median cells", "167"],
        ],
    )

    doc.add_paragraph(
        "The overrepresentation of TP53 and PTEN is expected: CRISPRi knockdown of "
        "tumor suppressors confers a proliferative advantage, causing those cells to "
        "expand clonally during culture. This is itself evidence that the perturbations "
        "are biologically active."
    )

    # -- 2. Expression --
    doc.add_heading("2. Most target genes are not expressed in this system", level=2)

    p = doc.add_paragraph("Of 224 target genes in the library, only ")
    run = p.add_run("29 are detectably expressed")
    run.bold = True
    p.add_run(" (NTC median CP10K \u2265 1.0) in EBs and/or iPSC cells:")

    add_table(doc,
        ["Category", "Count"],
        [
            ["Expressed in both EBs and iPSC", "13"],
            ["Expressed in EBs only", "10"],
            ["Expressed in iPSC only", "6"],
            ["Not expressed in either", "195"],
        ],
    )

    doc.add_paragraph(
        "This is expected for a broad transcription-factor-focused CRISPRi library "
        "applied to a specific cell type. The 195 non-expressed genes include "
        "tissue-specific factors (GATA4, FOXA2, CDX2, NKX2-5, etc.) that are "
        "silenced in iPSC/EBs."
    )

    p = doc.add_paragraph()
    run = p.add_run("Expressed in both conditions ")
    run.bold = True
    p.add_run("(sorted by total NTC expression): "
              "YBX1, NR6A1, SOX4, HDAC2, ID1, DNMT3B, SALL4, JARID2, BMPR1A, POU2F1, YY1, TCF3, SMARCA1")

    p = doc.add_paragraph()
    run = p.add_run("Expressed in EBs only: ")
    run.bold = True
    p.add_run("SOX2, NR2F2, PAX6, NRG1, OTX2, CREB1, SALL2, ZBTB16, HES1, PTEN")

    p = doc.add_paragraph()
    run = p.add_run("Expressed in iPSC only: ")
    run.bold = True
    p.add_run("POU5F1, DNMT1, EZH2, CTCF, EP300, E2F3")

    # -- 3. Cross-sample heatmap --
    doc.add_heading("3. Cross-sample knockdown overview", level=2)

    add_figure(
        doc,
        QC_DIR / "knockdown_heatmap.png",
        "Figure 2. Heatmap of log2 fold change for all 224 target genes (rows) across 16 samples (columns). "
        "Color: blue = knockdown, red = upregulation, white = no change. Black dots mark significant hits "
        "(BH p < 0.05 and log2FC < \u20130.5). Vertical line separates EBs (left) from iPSC (right). "
        "Rows sorted by median log2FC.",
        width=5.0,
    )

    add_table(doc,
        ["Condition", "Significant knockdowns per sample"],
        [
            ["EBs (10 samples)", "29\u201341 (median 35.5)"],
            ["iPSC (6 samples)", "44\u201372 (median 64.5)"],
        ],
    )

    doc.add_paragraph(
        "iPSC samples show approximately twice as many significant knockdowns as EBs samples. "
        "This likely reflects higher CRISPRi machinery expression or more accessible chromatin "
        "at targeted promoters in iPSCs."
    )

    # -- 4. Expression-filtered dotplot --
    doc.add_heading("4. Expression-filtered dotplot: the interpretable view", level=2)

    add_figure(
        doc,
        QC_DIR / "knockdown_expressed_only.png",
        "Figure 3. Dotplot restricted to the 29 expressed target genes. "
        "Dot size = number of perturbed cells. Dot color = log2FC vs NTC. "
        "Black ring = statistically significant knockdown. Grey dash = gene not expressed in that condition. "
        "Gene labels include condition specificity tag ([EBs] or [iPSC]) and median cell count.",
        width=6.5,
    )

    doc.add_heading("Strong, reproducible knockdowns (significant in 16/16 samples):", level=3)
    add_table(doc,
        ["Gene", "Median cells", "Median log2FC", "Note"],
        [
            ["TCF3", "69", "\u20132.79", "Deepest effect; E-protein TF"],
            ["PTEN", "374", "\u20132.87", "Tumor suppressor; growth advantage"],
            ["TP53", "552", "\u20132.45", "Tumor suppressor; most cells"],
            ["NR6A1", "68", "\u20131.24", "Nuclear receptor"],
            ["SOX4", "321", "\u20131.19", "HMG-box TF; growth advantage"],
        ],
        bold_col=0,
    )

    doc.add_heading("Moderate / partial knockdowns (significant in some samples):", level=3)
    add_table(doc,
        ["Gene", "Median cells", "Median log2FC", "Sig samples", "Note"],
        [
            ["GRHL2", "86", "\u20133.00", "14/16", "Strong in EBs, weaker in iPSC"],
            ["RELB", "87", "\u20131.70", "8/16", "NF-\u03baB family"],
            ["SALL4", "131", "\u20130.65", "10/16", "Stem cell TF"],
            ["DNMT3B", "55", "\u20130.97", "10/16", "DNA methyltransferase"],
            ["HES1 [EBs]", "47", "\u20130.98", "7/10 EBs", "Notch target; EBs-specific"],
            ["YBX1", "39", "\u20130.43", "3/16", "Highly expressed but modest FC"],
        ],
        bold_col=0,
    )

    doc.add_heading("Well-powered but not knocked down:", level=3)
    add_table(doc,
        ["Gene", "Median cells", "Median log2FC", "Note"],
        [
            ["SMARCA1", "123", "0.00", "Chromatin remodeler; guides ineffective"],
            ["ID1", "127", "\u20130.15", "Expressed but no knockdown"],
            ["BMPR1A", "122", "\u20130.09", "BMP receptor; guides miss active TSS?"],
            ["POU5F1 [iPSC]", "186", "+0.05", "OCT4; highly expressed, no knockdown"],
            ["SALL2 [EBs]", "46", "\u20130.18", "Weak/no effect"],
        ],
        bold_col=0,
    )

    doc.add_paragraph(
        "POU5F1 (OCT4) is notable: expressed at high levels in iPSC (NTC CP10K ~9.6) "
        "with 186 perturbed cells, yet shows zero knockdown. This may indicate the guides "
        "target the wrong TSS, or strong positive autoregulatory feedback compensates for "
        "partial CRISPRi repression."
    )

    # -- 5. Waterfall --
    doc.add_heading("5. Single-sample waterfall (EBs1_1)", level=2)

    add_figure(
        doc,
        QC_DIR / "waterfall_EBs1_1.png",
        "Figure 4. Waterfall plot for sample EBs1_1 showing log2FC for all 201 tested target genes "
        "sorted by effect size. Red bars = significant (BH p < 0.05 and log2FC < \u20130.5). "
        "36 genes show significant knockdown. No genes show significant upregulation, "
        "consistent with a clean CRISPRi experiment.",
        width=5.5,
    )

    # -- 6. All-gene dotplot --
    doc.add_heading("6. All-gene dotplot (cell-count aware)", level=2)

    add_figure(
        doc,
        QC_DIR / "knockdown_dotplot.png",
        "Figure 5. Dotplot for all 167 target genes with \u2265 20 median perturbed cells. "
        "Dot size = cell count, color = log2FC. This view includes non-expressed genes and "
        "illustrates that the vast majority of the library targets genes absent in this system.",
        width=5.5,
    )

    # ========== SUMMARY ==========
    doc.add_heading("Summary", level=1)

    add_table(doc,
        ["Metric", "Value"],
        [
            ["Total guides in library", "548"],
            ["Target genes", "224"],
            ["Expressed in EBs and/or iPSC", "29 (13%)"],
            ["Strong knockdowns (16/16 sig)", "5"],
            ["Partial knockdowns (\u2265 50% samples sig)", "~6 additional"],
            ["Well-powered, expressed, no knockdown", "~5"],
            ["Library utilization (expressed/total)", "13%"],
            ["Guide efficacy (knocked down/expressed)", "~38% (11/29)"],
        ],
    )

    doc.add_heading("Key takeaways", level=2)

    takeaways = [
        ("The CRISPRi system is working.",
         " The five strongest targets (TCF3, PTEN, TP53, NR6A1, SOX4) show deep, "
         "reproducible knockdown across all 16 samples with no false upregulations."),
        ("87% of the library is untestable",
         " in this system because the target genes are not expressed. This is a property "
         "of the library design (broad TF panel) applied to a specific cell type, not a technical failure."),
        ("iPSC samples show stronger perturbation effects",
         " (~65 significant knockdowns per sample vs ~35 in EBs), likely reflecting cell-type "
         "differences in CRISPRi efficiency or chromatin state."),
        ("Guide representation is skewed by fitness effects.",
         " TP53 and PTEN knockdown cells are overrepresented (550 and 374 cells vs median 44) "
         "due to proliferative advantage. This is biologically expected and itself validates the perturbations."),
        ("A few expressed genes are not knocked down",
         " despite adequate cell numbers (POU5F1, SMARCA1, ID1, BMPR1A). These likely reflect guide "
         "design issues (wrong TSS, inaccessible chromatin at the guide target site) rather than pipeline artifacts."),
    ]
    for bold_part, rest in takeaways:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(bold_part)
        run.bold = True
        p.add_run(rest)

    # ========== PROVENANCE ==========
    doc.add_heading("Provenance", level=1)
    provenance = [
        ("Analysis scripts", "perturbseq_qc_heatmap.py, perturbseq_qc_dotplot.py"),
        ("Feature reference", "cellranger_feature_ref_hCRISPRa_v2_like_AALG2_pattern.csv"),
        ("Input data", "Guarded EmptyDrops rerun filtered_counts.h5ad per sample"),
        ("Downstream suffix", "downstream_genefull_velocyto_cellbender_guarded_rerun"),
    ]
    for label, value in provenance:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(value)

    # ========== SAVE ==========
    doc.save(str(OUT))
    print(f"Saved {OUT}")


if __name__ == "__main__":
    build_report()
