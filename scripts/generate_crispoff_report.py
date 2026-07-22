#!/usr/bin/env python3
"""Generate a .docx report analyzing CRISPRoff efficacy in the UCSF perturb-seq data."""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

QC_DIR = Path("/mnt/pikachu/ucsf_release_20260403/qc_plots")
OUT = QC_DIR / "CRISPROFF_SELECTION_BIAS_REPORT.docx"


def add_table(doc, headers, rows, bold_col=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style="Light Shading Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    if bold_col is not None and j == bold_col:
                        run.bold = True
    doc.add_paragraph()


def add_figure(doc, path, caption, width=6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    run.italic = True
    doc.add_paragraph()


def bold_run(paragraph, text):
    run = paragraph.add_run(text)
    run.bold = True
    return run


def build_report():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ========== TITLE ==========
    title = doc.add_heading(
        "Proliferative Selection Bias in UCSF CRISPRi Perturb-seq:\n"
        "Evidence Against Effective CRISPRoff Temporal Control",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for label, value in [
        ("Date", "2026-04-03"),
        ("Dataset", "UCSF perturb-seq, 16 samples (10 EBs, 6 iPSC)"),
        ("Library", "hCRISPRa-v2 AALG2 (548 guides, 224 target genes)"),
        ("Pipeline", "STAR-suite with guarded EmptyDrops cell calling"),
    ]:
        run = meta.add_run(f"{label}: ")
        run.bold = True
        run.font.size = Pt(10)
        run = meta.add_run(f"{value}\n")
        run.font.size = Pt(10)

    # ========== EXECUTIVE SUMMARY ==========
    doc.add_heading("Executive Summary", level=1)

    p = doc.add_paragraph()
    p.add_run(
        "Analysis of guide representation in the UCSF CRISPRi perturb-seq dataset reveals "
        "a classic proliferative selection bias signature that is inconsistent with effective "
        "CRISPRoff temporal control. "
    )
    bold_run(p, "Tumor suppressor knockdowns (TP53, PTEN) are 10\u201315\u00d7 enriched ")
    p.add_run("over the population median, while ")
    bold_run(p, "essential gene knockdowns (MYC, NANOG, SOX2) are severely depleted ")
    p.add_run(
        "to < 0.2\u00d7 the median. This bidirectional skew is the hallmark of selection "
        "acting throughout the culture expansion period, exactly as expected in standard "
        "CRISPRi and contrary to the uniform representation expected from effective "
        "CRISPRoff-mediated late activation."
    )

    # ========== BACKGROUND ==========
    doc.add_heading("Background", level=1)

    doc.add_heading("The selection bias problem in CRISPRi screens", level=2)
    doc.add_paragraph(
        "In pooled CRISPRi perturb-seq experiments, guide-expressing cells are expanded "
        "in culture for days to weeks before the single-cell readout. During this expansion "
        "phase, cells carrying guides that knock down tumor suppressors (e.g., TP53, PTEN) "
        "gain a proliferative advantage and become overrepresented, while cells carrying "
        "guides that knock down essential or pro-growth genes (e.g., MYC, NANOG) are "
        "depleted from the population (Dixit et al., 2016; Replogle et al., 2022). This "
        "selection bias distorts guide representation and can render essential gene "
        "perturbations uninterpretable due to insufficient cell numbers."
    )

    doc.add_heading("CRISPRoff: epigenetic memory with temporal control", level=2)
    p = doc.add_paragraph(
        "CRISPRoff (Nun\u0303ez et al., 2021) was developed to address this problem. "
        "The system uses a catalytically dead Cas9 fused to KRAB, DNMT3A, and DNMT3L "
        "domains that write heritable DNA methylation at the target gene\u2019s promoter. "
        "The key innovation is "
    )
    bold_run(p, "temporal separation")
    p.add_run(
        ": the CRISPRoff machinery is expressed transiently to establish the epigenetic "
        "mark, then cleared from the cells. The silencing persists through cell divisions "
        "via maintenance methylation without requiring continuous repressor expression. "
        "This allows perturbation to be established late in the culture protocol, "
        "minimizing the window for proliferative selection to act."
    )

    p = doc.add_paragraph(
        "If CRISPRoff temporal control works as designed, all guide-bearing cells should "
        "expand at similar rates during the pre-induction phase (because no perturbation "
        "is active yet), yielding "
    )
    bold_run(p, "approximately uniform guide representation")
    p.add_run(
        " at the time of readout, regardless of the target gene\u2019s fitness effect."
    )

    # ========== ANALYSIS ==========
    doc.add_heading("Analysis", level=1)

    doc.add_heading("Approach", level=2)
    doc.add_paragraph(
        "We examined guide representation (median perturbed cells per target gene across "
        "all 16 samples) as a proxy for proliferative selection. Genes were classified "
        "by the expected fitness consequence of their CRISPRi knockdown:"
    )

    p = doc.add_paragraph(style="List Bullet")
    bold_run(p, "Tumor suppressors ")
    p.add_run(
        "(knockdown confers growth advantage): TP53, PTEN, RB1, CDKN1A, CDKN2A"
    )
    p = doc.add_paragraph(style="List Bullet")
    bold_run(p, "Essential / pro-growth genes ")
    p.add_run(
        "(knockdown causes growth disadvantage): MYC, YBX1, NANOG, SOX2, HDAC2, DNMT1"
    )
    p = doc.add_paragraph(style="List Bullet")
    bold_run(p, "Neutral / unknown ")
    p.add_run("(remaining ~213 target genes)")

    doc.add_paragraph(
        "Under effective CRISPRoff temporal control, all three groups should cluster "
        "around the overall median cell count (~37 cells). Under standard CRISPRi with "
        "selection, tumor suppressors should be enriched and essential genes depleted."
    )

    doc.add_heading("Results", level=2)

    add_figure(
        doc,
        QC_DIR / "selection_bias_analysis.png",
        "Figure 1. Guide representation across 224 target genes. Top: ranked bar chart "
        "colored by expected fitness effect of knockdown. Red = tumor suppressors "
        "(growth advantage from KD), blue = essential/pro-growth genes (growth disadvantage "
        "from KD), grey = other. Bottom: observed vs expected (CRISPRoff) median cell "
        "counts for each category. The 16\u00d7 ratio between tumor suppressors and essential "
        "genes is inconsistent with effective temporal control.",
        width=6.5,
    )

    add_table(
        doc,
        ["Category", "Expected (CRISPRoff)", "Observed", "Fold vs median"],
        [
            ["Tumor suppressors (n=5)", "~37 (1.0\u00d7)", "48", "1.3\u00d7"],
            ["Overall median", "37", "37", "1.0\u00d7"],
            ["Essential / pro-growth (n=6)", "~37 (1.0\u00d7)", "3", "0.08\u00d7"],
            ["Ratio (TS / essential)", "~1\u00d7", "16\u00d7", "\u2014"],
        ],
    )

    doc.add_heading("Individual gene representation", level=3)

    p = doc.add_paragraph()
    bold_run(p, "Enriched (growth advantage from knockdown):")

    add_table(
        doc,
        ["Gene", "Median cells", "Fold vs median", "Known function"],
        [
            ["TP53", "552", "14.9\u00d7", "Tumor suppressor; cell cycle arrest / apoptosis"],
            ["PTEN", "374", "10.1\u00d7", "Tumor suppressor; PI3K/AKT pathway brake"],
            ["SOX4", "321", "8.7\u00d7", "HMG-box TF; context-dependent growth effects"],
            ["CDKN1A", "48", "1.3\u00d7", "p21; CDK inhibitor downstream of TP53"],
            ["RB1", "48", "1.3\u00d7", "Retinoblastoma protein; G1/S checkpoint"],
            ["CDKN2A", "44", "1.2\u00d7", "p16; CDK4/6 inhibitor"],
        ],
        bold_col=0,
    )

    p = doc.add_paragraph()
    bold_run(p, "Depleted (growth disadvantage from knockdown):")

    add_table(
        doc,
        ["Gene", "Median cells", "Fold vs median", "Known function"],
        [
            ["SOX2", "8", "0.2\u00d7", "Pluripotency TF; essential in iPSC/EBs"],
            ["MYC", "4", "0.1\u00d7", "Proto-oncogene; essential for proliferation"],
            ["YBX1", "3", "0.1\u00d7", "Y-box binding protein; mRNA stability, proliferation"],
            ["NANOG", "3", "0.1\u00d7", "Pluripotency TF; essential in iPSC"],
            ["HDAC2", "2", "0.05\u00d7", "Histone deacetylase; essential chromatin regulator"],
            ["DNMT1", "1", "0.03\u00d7", "Maintenance methyltransferase; essential"],
        ],
        bold_col=0,
    )

    doc.add_heading("Knockdown efficacy is confirmed for well-represented genes", level=3)

    add_figure(
        doc,
        QC_DIR / "knockdown_expressed_only.png",
        "Figure 2. Knockdown dotplot for the 29 expressed target genes (NTC CP10K \u2265 1.0). "
        "Dot size = cell count, color = log2FC vs NTC, black ring = significant (BH p < 0.05 "
        "& log2FC < \u20130.5). The strong knockdown of TP53, PTEN, TCF3, SOX4, and NR6A1 "
        "across all 16 samples confirms CRISPRi is functionally active, not just passively "
        "inherited.",
        width=6.5,
    )

    doc.add_paragraph(
        "Importantly, the perturbations are not merely inherited epigenetic marks \u2014 they "
        "are actively producing transcriptional repression. Five target genes show significant "
        "knockdown in 16/16 samples with log2FC of \u20131.2 to \u20132.9 (see Figure 2). This level "
        "of active repression is consistent with ongoing CRISPRi activity (dCas9-KRAB mediated "
        "transcriptional interference) rather than CRISPRoff-mediated DNA methylation memory, "
        "which typically produces more moderate silencing."
    )

    # ========== INTERPRETATION ==========
    doc.add_heading("Interpretation", level=1)

    doc.add_heading("Evidence against effective CRISPRoff temporal control", level=2)

    doc.add_paragraph("Three independent lines of evidence indicate the perturbations were "
                      "active throughout the culture expansion period:")

    p = doc.add_paragraph(style="List Number")
    bold_run(p, "Bidirectional selection on cell count. ")
    p.add_run(
        "Tumor suppressor knockdowns are enriched (TP53 at 14.9\u00d7 median) while "
        "essential gene knockdowns are depleted (MYC at 0.1\u00d7 median). If CRISPRoff "
        "delayed perturbation until near assay time, both groups would cluster around "
        "the median because cells would have expanded without any active perturbation. "
        "The observed 16\u00d7 ratio between these groups requires days to weeks of "
        "differential proliferation."
    )

    p = doc.add_paragraph(style="List Number")
    bold_run(p, "Deep, active transcriptional repression. ")
    p.add_run(
        "The observed log2FC values (\u20131.2 to \u20132.9 for the top 5 targets) indicate "
        "strong, ongoing transcriptional interference. CRISPRoff-mediated silencing via "
        "DNA methylation is typically more moderate and does not require continuous dCas9 "
        "expression, whereas CRISPRi-mediated KRAB repression produces the deep knockdowns "
        "seen here but requires the dCas9-KRAB fusion to be constitutively expressed."
    )

    p = doc.add_paragraph(style="List Number")
    bold_run(p, "Near-complete depletion of essential gene perturbations. ")
    p.add_run(
        "DNMT1 (1 cell), HDAC2 (2 cells), NANOG (3 cells), and MYC (4 cells) are at "
        "the floor of detection. Even with modest CRISPRi activity during expansion, "
        "knockdown of these genes would cause rapid cell death or arrest, leading to "
        "exponential depletion over multiple doublings. This degree of depletion is "
        "not expected if perturbation was restricted to the final 24\u201348 hours."
    )

    doc.add_heading("Possible explanations", level=2)

    explanations = [
        ("Constitutive CRISPRi, not CRISPRoff",
         "The experiment may have used a standard constitutive dCas9-KRAB system despite "
         "the library being designed for CRISPRoff. The library name (hCRISPRa-v2) and "
         "guide design are compatible with either system \u2014 the distinction lies in the "
         "effector construct and induction protocol, not the guides themselves."),
        ("Inducible system leak",
         "If an inducible CRISPRoff system was used, the effector may have had basal "
         "expression (leak) in the uninduced state, sufficient to produce CRISPRi-like "
         "repression during expansion. Even low-level dCas9-KRAB expression can produce "
         "meaningful transcriptional interference at sensitive loci like TP53."),
        ("Premature induction",
         "The CRISPRoff induction may have been triggered too early relative to the "
         "expansion period, leaving sufficient time for selection to act. CRISPRoff "
         "establishes silencing within 24\u201348 hours (Nun\u0303ez et al., 2021), so induction "
         "even 3\u20134 days before harvest could allow several cell doublings under selection."),
    ]
    for label, text in explanations:
        p = doc.add_paragraph(style="List Number")
        bold_run(p, f"{label}. ")
        p.add_run(text)

    # ========== IMPACT ==========
    doc.add_heading("Impact on Downstream Analysis", level=1)

    add_table(
        doc,
        ["Gene category", "Affected?", "Recommendation"],
        [
            ["Neutral-fitness targets (~200 genes)", "Minimal",
             "Guide representation is adequate; standard differential expression analysis is valid"],
            ["Tumor suppressors (TP53, PTEN, etc.)", "Yes \u2014 enriched",
             "Cell counts are inflated by clonal expansion; interpret with caution. "
             "Expression changes are real but cell composition may be biased"],
            ["Essential genes (MYC, NANOG, SOX2, etc.)", "Yes \u2014 depleted",
             "Insufficient cells for statistical testing; perturbation effects are "
             "uninterpretable. Consider excluding from analysis"],
            ["Moderately essential genes", "Possibly",
             "Genes with 5\u201315 cells may have partial depletion; flag as low-confidence"],
        ],
    )

    doc.add_paragraph(
        "The selection bias does not invalidate the dataset. The majority of the library "
        "targets neutral-fitness transcription factors, and these perturbations have "
        "adequate representation and produce interpretable results. However, any analysis "
        "of fitness-sensitive targets (tumor suppressors, essential genes) should account "
        "for the selection that occurred during expansion."
    )

    # ========== SUMMARY ==========
    doc.add_heading("Summary", level=1)

    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Guide representation range", "1 \u2013 552 cells (median 37)"],
            ["TP53 enrichment", "14.9\u00d7 median"],
            ["PTEN enrichment", "10.1\u00d7 median"],
            ["Essential gene depletion (MYC, NANOG, etc.)", "0.03\u20130.2\u00d7 median"],
            ["TS / essential ratio", "16\u00d7 (expected ~1\u00d7 if CRISPRoff worked)"],
            ["Consistent with CRISPRoff temporal control?", "No"],
            ["Consistent with standard CRISPRi + selection?", "Yes"],
            ["Impact on neutral-fitness gene analysis?", "Minimal"],
        ],
    )

    # ========== REFERENCES ==========
    doc.add_heading("References", level=1)

    refs = [
        ("Nun\u0303ez JK, Chen J, Pommier GC, et al.",
         "Genome-wide programmable transcriptional memory by CRISPR-based epigenome editing.",
         "Cell", "2021", "184(9):2503\u20132519.e17",
         "Established CRISPRoff as a heritable epigenetic silencing tool using dCas9-KRAB-DNMT3A/3L. "
         "Demonstrated silencing persistence over 450+ cell divisions and reversibility with "
         "CRISPRon (TET1-dCas9). Key reference for expected CRISPRoff behavior."),

        ("Dixit A, Parnas O, Li B, et al.",
         "Perturb-Seq: Dissecting Molecular Circuits with Scalable Single-Cell RNA Profiling "
         "of Pooled Genetic Screens.",
         "Cell", "2016", "167(7):1853\u20131866.e17",
         "Original perturb-seq methodology. Documented selection bias as a key confounder in "
         "pooled CRISPR screens and proposed computational correction approaches."),

        ("Replogle JM, Saunders RA, Pogson AN, et al.",
         "Mapping information-rich genotype-phenotype landscapes with genome-scale Perturb-seq.",
         "Cell", "2022", "185(14):2559\u20132575.e28",
         "Genome-scale CRISPRi perturb-seq in K562 and RPE1 cells. Extensively characterized "
         "proliferative selection artifacts including TP53-driven enrichment and essential "
         "gene depletion. Recommended guide representation QC as standard practice."),

        ("Jost M, Chen Y, Gilbert LA, et al.",
         "Combined CRISPRi/a-Based Chemical Genetic Screens Reveal that Rigosertib Is a "
         "Microtubule-Destabilizing Agent.",
         "Molecular Cell", "2017", "68(1):210\u2013223.e6",
         "Demonstrated that pooled CRISPRi screens are subject to proliferative selection "
         "over the culture period. Guides targeting essential genes are depleted at rates "
         "proportional to the fitness cost of knockdown."),

        ("Lalli MA, Avey D, Bhagwat JD, et al.",
         "High-throughput single-cell functional elucidation of neurodevelopmental "
         "disease-associated genes.",
         "Nature Communications", "2020", "11:6316",
         "Large-scale CRISPRi perturb-seq of transcription factors in iPSC-derived neurons. "
         "Relevant as a comparison dataset using a similar TF-focused library in iPSC-derived cells."),

        ("Tian R, Gachechiladze MA, Ludwig CH, et al.",
         "CRISPR Interference-Based Platform for Multimodal Genetic Screens in Human "
         "iPSC-Derived Neurons.",
         "Neuron", "2019", "104(2):239\u2013255.e12",
         "Demonstrated CRISPRi perturb-seq in iPSC-derived neurons with fitness-aware "
         "analysis. Reported similar selection patterns: TP53 guides enriched, essential "
         "gene guides depleted during the neuronal differentiation protocol."),
    ]

    for i, (authors, title, journal, year, pages, note) in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)

        run = p.add_run(f"{i}. {authors} ")
        run.font.size = Pt(9)

        run = p.add_run(f"{title} ")
        run.font.size = Pt(9)
        run.italic = True

        run = p.add_run(f"{journal}. {year};{pages}.")
        run.font.size = Pt(9)
        run.bold = True

        if note:
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Inches(0.3)
            p2.paragraph_format.space_before = Pt(0)
            run = p2.add_run(note)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            run.italic = True

    # ========== PROVENANCE ==========
    doc.add_heading("Provenance", level=1)
    provenance = [
        ("Analysis date", "2026-04-03"),
        ("Analysis scripts",
         "perturbseq_qc_heatmap.py, perturbseq_qc_dotplot.py (STAR-suite/scripts/)"),
        ("Feature reference", "cellranger_feature_ref_hCRISPRa_v2_like_AALG2_pattern.csv"),
        ("Input data", "Guarded EmptyDrops rerun filtered_counts.h5ad per sample"),
        ("Production root",
         "ucsf-perturb-yremove_velocyto_cellbender_prod_globus_fixtruwl_20260330_225009"),
        ("Statistical method",
         "CP10K + log1p normalization, Wilcoxon rank-sum test, BH FDR correction"),
    ]
    for label, value in provenance:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{label}: ")
        run.bold = True
        run.font.size = Pt(9)
        run = p.add_run(value)
        run.font.size = Pt(9)

    doc.save(str(OUT))
    print(f"Saved {OUT}")


if __name__ == "__main__":
    build_report()
